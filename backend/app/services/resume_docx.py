from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor

from app.services.resume_content import ResumeItem, build_resume_content

_INK = RGBColor(0x1A, 0x1A, 0x1A)
_MUTED = RGBColor(0x33, 0x33, 0x33)


def generate_resume_docx() -> bytes:
    c = build_resume_content()
    doc = Document()

    core = doc.core_properties
    core.title = f"{c.name} - Resume"
    core.author = c.name
    core.subject = c.title
    core.keywords = c.keywords

    section = doc.sections[0]
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    section.top_margin = Mm(13)
    section.bottom_margin = Mm(13)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(1)

    def para(text: str = "", bold_prefix: str = "", size: int = 10,
             bold: bool = False, color: RGBColor = _INK, indent: float = 0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        if indent:
            p.paragraph_format.left_indent = Pt(indent)
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            r.font.size = Pt(size)
            r.font.color.rgb = color
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(size)
            r.font.color.rgb = color
        return p

    def heading(title: str):
        p = para(title, size=11, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        # Thin bottom border under each section heading (ATS-safe, no tables).
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "1A1A1A")
        borders.append(bottom)
        pPr.append(borders)

    def items(title: str, entries: list[ResumeItem]):
        if not entries:
            return
        heading(title)
        for it in entries:
            para(f" ({it.year}): {it.description}", bold_prefix=f"- {it.title}", indent=10)

    # ---- Header ----
    name_p = para(c.name, size=17, bold=True)
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para(c.title, size=9, color=_MUTED)
    para(c.contact_line, size=9, color=_MUTED)
    if c.links:
        para(" | ".join(c.links), size=9, color=_MUTED)

    heading("PROFESSIONAL SUMMARY")
    para(c.summary)

    if c.skills:
        heading("TECHNICAL SKILLS")
        for cat, names in c.skills:
            para(f" {names}", bold_prefix=f"{cat}:")

    if c.experiences:
        heading("PROFESSIONAL EXPERIENCE")
        for exp in c.experiences:
            para(f", {exp.company} | {exp.location} | {exp.period}", bold_prefix=exp.role)
            for h in exp.bullets:
                para(f"- {h}", indent=10)
            if exp.tech:
                p = para(f" {', '.join(exp.tech)}", bold_prefix="Technologies:",
                         color=_MUTED, indent=10)
                p.paragraph_format.space_after = Pt(4)

    if c.projects:
        heading("PROJECTS")
        for proj in c.projects:
            para(f" | {', '.join(proj.tech)}", bold_prefix=proj.title)
            para(f"- {proj.description}", indent=10)

    if c.education:
        heading("EDUCATION")
        for ed in c.education:
            para(f", {ed.institution} | {ed.period}", bold_prefix=ed.degree)

    items("ACHIEVEMENTS & AWARDS", c.awards)
    items("CERTIFICATIONS", c.certifications)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
